from flask import Flask, render_template, request, send_file
<<<<<<< HEAD
import os # Deze import is nu toegevoegd of gecontroleerd
=======
import os
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
from io import BytesIO
from PIL import Image, ExifTags, ImageOps
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
<<<<<<< HEAD
# De UPLOAD_FOLDER wordt in deze applicatie niet gebruikt om bestanden fysiek op te slaan,
# omdat ze direct in het geheugen worden verwerkt. Maar de map wordt wel aangemaakt.
app.config['UPLOAD_FOLDER'] = 'uploads' 
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Breedte van de inhoud van het document in centimeters
DOCUMENT_CONTENT_WIDTH_CM = 17.0
# Conversiefactor van centimeters naar inches
CM_PER_INCH = 2.54

# Instellingen voor beeldkwaliteit en resampling
IMAGE_SETTINGS = {
    'original': {'resize': False, 'quality': 95},
    'hd': {'ppi': 330, 'quality': 90},
    'print': {'ppi': 220, 'quality': 85}, # Standaard instelling
=======
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

DOCUMENT_CONTENT_WIDTH_CM = 17.0
CM_PER_INCH = 2.54

IMAGE_SETTINGS = {
    'original': {'resize': False, 'quality': 95},
    'hd': {'ppi': 330, 'quality': 90},
    'print': {'ppi': 220, 'quality': 85}, # Standaard
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
    'web': {'ppi': 150, 'quality': 80},
    'email': {'ppi': 96, 'quality': 75},
}

def fix_image_orientation(img):
    """
    Roteert de afbeelding fysiek op basis van de EXIF-oriëntatie tag en verwijdert de tag.
<<<<<<< HEAD
    Dit zorgt ervoor dat foto's die bijvoorbeeld met een telefoon zijn gemaakt
    correct worden weergegeven, ongeacht hun originele oriëntatie.
    Gebruikt ImageOps.exif_transpose() voor een robuuste oplossing.

    Args:
        img (PIL.Image.Image): Het Pillow Image object.
    Returns:
        PIL.Image.Image: De geroteerde afbeelding met gecorrigeerde oriëntatie.
=======
    Gebruikt ImageOps.exif_transpose() voor een robuuste oplossing.
    Args:
        img (PIL.Image.Image): De Pillow Image object.
    Returns:
        PIL.Image.Image: De geroteerde afbeelding.
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
    """
    try:
        img = ImageOps.exif_transpose(img)
    except Exception as e:
<<<<<<< HEAD
        # Log een fout als het corrigeren van de oriëntatie mislukt,
        # maar ga verder met de verwerking van de afbeelding.
=======
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
        print(f"Kon EXIF-oriëntatie niet corrigeren met exif_transpose: {e}")
    return img

@app.route('/')
def index():
    """
<<<<<<< HEAD
    Rendert de hoofdpagina van de web-app (index.html),
    waar gebruikers foto's kunnen uploaden en instellingen kunnen kiezen.
=======
    Renders the main upload page.
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
    """
    return render_template('index.html')

@app.route('/generate-document', methods=['POST'])
def generate_document():
    """
<<<<<<< HEAD
    Verwerkt de ingediende gegevens van het formulier:
    - Haalt de geüploade foto's op.
    - Past afbeeldingskwaliteit en grootte aan.
    - Genereert een Word-document met de foto's in een tabelindeling.
    - Stuurt het gegenereerde Word-document terug als download.
    """
    # Controleer of er een 'photos' veld in de request zit
    if 'photos' not in request.files:
        return "Geen foto's gevonden in het verzoek", 400

    photos = request.files.getlist('photos')
    # Controleer of er daadwerkelijk bestanden zijn geselecteerd
    if not photos or all(p.filename == '' for p in photos):
        return "Geen bestanden geselecteerd", 400

    # Haal de formulierinstellingen op
    quality_setting = request.form.get('image_quality', 'print')
    num_columns = int(request.form.get('num_columns', 2))
    whitespace_mm = float(request.form.get('whitespace_mm', 5))
    include_captions = request.form.get('include_captions') == 'on'

    # Haal de bijschriften op en verwijder witruimte
    captions_raw = request.form.getlist('caption[]')
    captions = [c.strip() for c in captions_raw]

    # Berekeningen voor de tabelindeling
    column_width_cm = DOCUMENT_CONTENT_WIDTH_CM / num_columns
    padding_cm = whitespace_mm / 10.0 # Millimeters converteren naar centimeters

    effective_image_width_in_cell_cm = column_width_cm - (2 * padding_cm)
    # Zorg ervoor dat de effectieve breedte van de afbeelding niet negatief of nul wordt.
    # Als dit gebeurt, betekent het dat de padding te groot is voor de kolombreedte.
    if effective_image_width_in_cell_cm <= 0:
        # Overweeg hier een betere foutmelding aan de gebruiker terug te geven
        # in plaats van stilzwijgend een minimale breedte te forceren.
        effective_image_width_in_cell_cm = 1.0 
=======
    Handles the form submission, processes images, and generates the Word document.
    """
    if 'photos' not in request.files:
        return "No photo part in the request", 400

    photos = request.files.getlist('photos')
    if not photos or all(p.filename == '' for p in photos):
        return "No selected file", 400

    quality_setting = request.form.get('image_quality', 'print')
    num_columns = int(request.form.get('num_columns', 2)) # Standaard 2 kolommen
    whitespace_mm = float(request.form.get('whitespace_mm', 5))
    include_captions = request.form.get('include_captions') == 'on'

    captions_raw = request.form.getlist('caption[]')
    captions = [c.strip() for c in captions_raw]

    column_width_cm = DOCUMENT_CONTENT_WIDTH_CM / num_columns
    padding_cm = whitespace_mm / 10.0

    effective_image_width_in_cell_cm = column_width_cm - (2 * padding_cm)
    if effective_image_width_in_cell_cm <= 0:
        effective_image_width_in_cell_cm = 1.0
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe

    # --- Start Word document generatie ---
    document = Document()

    # Stel paginamarges in
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    document.add_heading('Gegenereerde Fotoreportage', level=1)

<<<<<<< HEAD
    # Bereken het aantal rijen dat nodig is voor de tabel
    num_rows = (len(photos) + num_columns - 1) // num_columns
    table = document.add_table(rows=num_rows, cols=num_columns)
    table.autofit = False # Schakel automatische aanpassing van kolombreedtes uit
    table.allow_autofit = False # Voorkomt dat Word de tabel automatisch aanpast

    # Stel de breedte van elke kolom in
=======
    num_rows = (len(photos) + num_columns - 1) // num_columns
    table = document.add_table(rows=num_rows, cols=num_columns)
    table.autofit = False
    table.allow_autofit = False
    
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
    for i in range(num_columns):
        table.columns[i].width = Cm(column_width_cm)

    row_index = 0
    col_index = 0

<<<<<<< HEAD
    # Loop door elke geüploade foto
    for i, photo_file in enumerate(photos):
        cell = table.cell(row_index, col_index)
        
        # Verwijder alle bestaande alinea's in de cel om een schone lei te garanderen
=======
    for i, photo_file in enumerate(photos):
        cell = table.cell(row_index, col_index)
        
        # Verwijder ALLE bestaande alinea's in de cel
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
        for paragraph in list(cell.paragraphs):
            if paragraph._element.getparent() is not None:
                cell._tc.remove(paragraph._element)

<<<<<<< HEAD
        # Stel de celmarges (padding) in met behulp van OxmlElement
        # Dit geeft fijne controle over de lay-out binnen de cel
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement('w:tcMar')
        
        # Converteer centimeters naar dxa (twintigste van een punt), een eenheid die Word gebruikt
        margin_dxa = str(int(padding_cm * 567)) 

        # Voeg marge-elementen toe voor alle vier de zijden
=======
        # Stel de celmarges in voor top, bottom, left, right
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement('w:tcMar')
        
        margin_dxa = str(int(padding_cm * 567)) 

>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
        top_margin = OxmlElement('w:top')
        top_margin.set(qn('w:w'), margin_dxa)
        top_margin.set(qn('w:type'), 'dxa')
        tc_mar.append(top_margin)

        bottom_margin = OxmlElement('w:bottom')
        bottom_margin.set(qn('w:w'), margin_dxa)
        bottom_margin.set(qn('w:type'), 'dxa')
        tc_mar.append(bottom_margin)
        
        left_margin = OxmlElement('w:left')
        left_margin.set(qn('w:w'), margin_dxa)
        left_margin.set(qn('w:type'), 'dxa')
        tc_mar.append(left_margin)

        right_margin = OxmlElement('w:right')
        right_margin.set(qn('w:w'), margin_dxa)
        right_margin.set(qn('w:type'), 'dxa')
        tc_mar.append(right_margin)

<<<<<<< HEAD
        tc_pr.append(tc_mar) # Voeg de marges toe aan de celproperties

        # Verwerk de foto als de bestandsnaam niet leeg is
        if photo_file.filename != '':
            try:
                # Open de afbeelding vanuit de geüploade stream
                img = Image.open(photo_file.stream)
                
                # Corrigeer de oriëntatie op basis van EXIF-gegevens
                img = fix_image_orientation(img)

                # Converteer afbeeldingen met transparantie (RGBA, P) naar RGB,
                # aangezien JPEG geen transparantie ondersteunt
=======
        tc_pr.append(tc_mar)


        if photo_file.filename != '':
            try:
                img = Image.open(photo_file.stream)
                
                img = fix_image_orientation(img)

>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
                if img.mode in ['RGBA', 'P']:
                    img = img.convert('RGB')

                target_pixel_width = img.width
                target_pixel_height = img.height
                quality_val = IMAGE_SETTINGS[quality_setting].get('quality', 85)

                if quality_setting != 'original':
<<<<<<< HEAD
                    # Schaal de afbeelding naar de gewenste PPI en effectieve breedte
=======
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
                    ppi = IMAGE_SETTINGS[quality_setting]['ppi']
                    calculated_target_width_px = int(effective_image_width_in_cell_cm * (ppi / CM_PER_INCH))
                    
                    if img.width > calculated_target_width_px:
<<<<<<< HEAD
                        # Verklein alleen als de originele breedte groter is dan de berekende doelbreedte
=======
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
                        target_pixel_width = calculated_target_width_px
                        aspect_ratio = img.height / img.width
                        target_pixel_height = int(target_pixel_width * aspect_ratio)
                        img = img.resize((target_pixel_width, target_pixel_height), Image.LANCZOS)
<<<<<<< HEAD
                    # Anders (afbeelding is al kleiner of gelijk aan de doelbreedte), behoud de originele grootte.
                else:
                    # Voor 'original' setting, behoud de originele grootte tenzij deze te groot is voor de cel
                    # We gebruiken hier een hoge PPI (330) als referentie voor de maximale pixelbreedte
=======
                    else:
                        target_pixel_width = img.width
                        target_pixel_height = img.height
                else:
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
                    max_cell_px_width_for_original = int(effective_image_width_in_cell_cm * (330 / CM_PER_INCH))
                    if img.width > max_cell_px_width_for_original:
                        target_pixel_width = max_cell_px_width_for_original
                        aspect_ratio = img.height / img.width
                        target_pixel_height = int(target_pixel_width * aspect_ratio)
                        img = img.resize((target_pixel_width, target_pixel_height), Image.LANCZOS)
<<<<<<< HEAD
                    # Anders, behoud de originele grootte.

                # Sla de verwerkte afbeelding op in een BytesIO-object
                # Dit houdt de afbeelding in het geheugen in plaats van op schijf
                img_byte_arr = BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=quality_val) 
                img_byte_arr.seek(0) # Zet de pointer terug naar het begin van de stream

                # Voeg de afbeelding toe aan de cel
                picture_paragraph = cell.add_paragraph() 
                run = picture_paragraph.add_run()
                run.add_picture(img_byte_arr, width=Cm(effective_image_width_in_cell_cm))
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centreer de afbeelding

                # Nul witruimte boven en onder de afbeelding-alinea voor strakke opmaak
=======
                    else:
                        target_pixel_width = img.width
                        target_pixel_height = img.height

                img_byte_arr = BytesIO()
                img.save(img_byte_arr, format='JPEG', quality=quality_val) 
                img_byte_arr.seek(0)

                # Creëer de alinea voor de foto NADAT de cel is geleegd
                picture_paragraph = cell.add_paragraph() 
                run = picture_paragraph.add_run()
                run.add_picture(img_byte_arr, width=Cm(effective_image_width_in_cell_cm))
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Nul witruimte boven en onder de afbeelding-alinea
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
                paragraph_format = picture_paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

<<<<<<< HEAD
                # Voeg bijschrift toe indien gewenst
                current_caption_text = captions[i] if include_captions and i < len(captions) else ''
                if include_captions:
                    caption_paragraph = cell.add_paragraph(current_caption_text)
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_paragraph.style = 'Caption' # Pas de 'Caption' stijl toe (moet bestaan in Word)
                    # Nul witruimte boven het bijschrift als het direct na een afbeelding komt
                    caption_paragraph_format = caption_paragraph.paragraph_format
                    caption_paragraph_format.space_before = Pt(0)
                    
            except Exception as e:
                # Vang fouten op die optreden tijdens de verwerking van een afbeelding
                print(f"Fout bij verwerken afbeelding {photo_file.filename}: {e}")
                error_paragraph = cell.add_paragraph(f"Fout bij laden afbeelding {photo_file.filename}")
                error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Lijn de inhoud van de cel altijd bovenaan uit voor een consistente lay-out
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                
        # Ga naar de volgende kolom of naar de volgende rij als de huidige rij vol is
=======
                current_caption_text = captions[i] if include_captions and i < len(captions) else ''
                
                if include_captions:
                    caption_paragraph = cell.add_paragraph(current_caption_text)
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_paragraph.style = 'Caption'
                    # Nul witruimte boven het bijschrift als het volgt op een afbeelding
                    caption_paragraph_format = caption_paragraph.paragraph_format
                    caption_paragraph_format.space_before = Pt(0)
                    
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            except Exception as e:
                print(f"Error processing image {photo_file.filename}: {e}")
                error_paragraph = cell.add_paragraph(f"Fout bij laden afbeelding {photo_file.filename}")
                error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
        col_index += 1
        if col_index >= num_columns:
            col_index = 0
            row_index += 1

<<<<<<< HEAD
    # Sla het gegenereerde Word-document op in een BytesIO-object
    document_stream = BytesIO()
    document.save(document_stream)
    document_stream.seek(0) # Belangrijk: zet de pointer terug naar het begin van de stream

    # Stuur het Word-document terug naar de gebruiker als download
=======
    document_stream = BytesIO()
    document.save(document_stream)
    document_stream.seek(0)

>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
    return send_file(
        document_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="Fotoreportage.docx"
    )

<<<<<<< HEAD
# Dit blok code wordt alleen uitgevoerd wanneer het script direct wordt aangeroepen (bijv. 'python app.py').
# In een productieomgeving (zoals Heroku of Render) zal Gunicorn je Flask-app anders starten
# en dit blok negeren.
if __name__ == '__main__':
    # De omgevingsvariabele 'PORT' wordt vaak ingesteld door hostingplatforms.
    # Als 'PORT' niet is ingesteld, gebruiken we standaard poort 5000 voor lokale ontwikkeling.
    port = int(os.environ.get("PORT", 5000))
    # 'host='0.0.0.0'' zorgt ervoor dat de ontwikkelserver luistert op alle netwerkinterfaces,
    # wat nuttig kan zijn in containeromgevingen of als je vanaf een andere machine toegang wilt krijgen.
    app.run(debug=True, host='0.0.0.0', port=port)

=======
if __name__ == '__main__':
    app.run(debug=True)
>>>>>>> 65cd5d93e3b7c8a988c0378e762c7cbe81c24efe
